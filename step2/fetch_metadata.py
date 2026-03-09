#!/usr/bin/env python3
"""
Step 2: Fetch Play Store metadata for each competitor and compile into a .docx
Reads blockchain_competitors.csv from ../step1/, outputs competitor_metadata.docx
"""

import csv
import glob
import os
import re
import sys
import time

from google_play_scraper import app as gplay_app
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── Constants ──────────────────────────────────────────────────────────────
STEP1_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "step1")
OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_DOCX = os.path.join(OUTPUT_DIR, "competitor_metadata.docx")
BLOCKCHAIN_PKG = "piuk.blockchain.android"


def find_csv():
    """Find the first .csv file in the step1 directory."""
    pattern = os.path.join(STEP1_DIR, "*.csv")
    files = sorted(glob.glob(pattern))
    if not files:
        print(f"ERROR: No CSV files found in {STEP1_DIR}")
        sys.exit(1)
    return files[0]


def strip_html(text):
    """Remove HTML tags and collapse whitespace."""
    if not text:
        return ""
    # Remove HTML tags
    text = re.sub(r'<[^>]+>', '', text)
    # Decode common HTML entities
    text = text.replace('&amp;', '&').replace('&lt;', '<').replace('&gt;', '>')
    text = text.replace('&nbsp;', ' ').replace('&quot;', '"').replace('&#39;', "'")
    # Collapse runs of 3+ newlines into 2 (one blank line max)
    text = re.sub(r'\n{3,}', '\n\n', text)
    # Remove leading/trailing blank lines
    text = text.strip()
    return text


def fetch_app_data(package_id):
    """Fetch app metadata from Google Play Store."""
    result = gplay_app(package_id, lang='en', country='us')

    # Build tags from genre + categories
    tags = set()
    if result.get('genre'):
        tags.add(result['genre'])
    for cat in (result.get('categories') or []):
        name = cat.get('name', '')
        if name:
            tags.add(name)

    return {
        'title': strip_html(result.get('title', package_id)),
        'summary': strip_html(result.get('summary', '')),
        'description': strip_html(result.get('description', '')),
        'developer': result.get('developer', ''),
        'tags': ', '.join(sorted(tags)) if tags else '',
    }


def build_docx(entries, total):
    """Create the formatted .docx document."""
    from docx.oxml.ns import qn
    from docx.shared import Twips

    doc = Document()

    # Set default font and margins
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    # Reduce default paragraph spacing to zero
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    for i, entry in enumerate(entries):
        if entry is None:
            continue

        # Title — bold, 14pt, dark blue, small space after
        title_para = doc.add_paragraph()
        title_para.paragraph_format.space_before = Pt(6) if i > 0 else Pt(0)
        title_para.paragraph_format.space_after = Pt(2)
        title_run = title_para.add_run(entry['title'])
        title_run.bold = True
        title_run.font.size = Pt(14)
        title_run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        title_run.font.name = 'Calibri'

        # Developer — label bold, value normal
        dev_para = doc.add_paragraph()
        dev_para.paragraph_format.space_after = Pt(2)
        dev_label = dev_para.add_run('Developer: ')
        dev_label.bold = True
        dev_label.font.size = Pt(11)
        dev_label.font.name = 'Calibri'
        dev_value = dev_para.add_run(entry.get('developer', ''))
        dev_value.font.size = Pt(11)
        dev_value.font.name = 'Calibri'

        # Tags — label bold, value normal
        tags_para = doc.add_paragraph()
        tags_para.paragraph_format.space_after = Pt(2)
        tags_label = tags_para.add_run('Tags: ')
        tags_label.bold = True
        tags_label.font.size = Pt(11)
        tags_label.font.name = 'Calibri'
        tags_value = tags_para.add_run(entry.get('tags', ''))
        tags_value.font.size = Pt(11)
        tags_value.font.name = 'Calibri'

        # Short Description — label bold, value normal, 11pt
        desc_para = doc.add_paragraph()
        desc_para.paragraph_format.space_after = Pt(4)
        label_run = desc_para.add_run('Short Description: ')
        label_run.bold = True
        label_run.font.size = Pt(11)
        label_run.font.name = 'Calibri'
        value_run = desc_para.add_run(entry['summary'])
        value_run.font.size = Pt(11)
        value_run.font.name = 'Calibri'

        # Full description — single paragraph, skip empty lines, soft breaks
        description = entry.get('description', '')
        if description:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            # Filter out empty/whitespace-only lines
            lines = [ln for ln in description.split('\n') if ln.strip()]
            for j, line in enumerate(lines):
                run = p.add_run(line)
                run.font.size = Pt(11)
                run.font.name = 'Calibri'
                if j < len(lines) - 1:
                    run.add_break()

        # Grey separator line — tight spacing
        separator_para = doc.add_paragraph()
        separator_para.paragraph_format.space_before = Pt(4)
        separator_para.paragraph_format.space_after = Pt(0)
        sep_run = separator_para.add_run('─' * 50)
        sep_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        sep_run.font.size = Pt(9)
        sep_run.font.name = 'Calibri'

    doc.save(OUTPUT_DOCX)


def main():
    print("=" * 60)
    print("Step 2: Fetch Play Store Metadata")
    print("=" * 60)

    # Find and read CSV
    csv_path = find_csv()
    print(f"\nDetected CSV: {os.path.basename(csv_path)}")
    print(f"Full path: {csv_path}\n")

    with open(csv_path, 'r', encoding='utf-8') as f:
        reader = csv.DictReader(f)
        rows = list(reader)

    package_ids = [row['play_store_package_id'] for row in rows]

    # Always include Blockchain.com first, then the rest (deduplicated)
    ordered_ids = [BLOCKCHAIN_PKG]
    for pid in package_ids:
        if pid != BLOCKCHAIN_PKG:
            ordered_ids.append(pid)

    total = len(ordered_ids)
    print(f"Total apps to process: {total}\n")

    # Fetch metadata for each app
    entries = []
    successes = 0
    failures = []

    for i, pkg_id in enumerate(ordered_ids, 1):
        try:
            data = fetch_app_data(pkg_id)
            entries.append(data)
            successes += 1
            print(f"✓ {i}/{total} — {data['title']}")
            # Small delay to avoid rate limiting
            time.sleep(0.5)
        except Exception as e:
            entries.append(None)
            failures.append((pkg_id, str(e)))
            print(f"✗ {i}/{total} — FAILED ({pkg_id}): {e}")

    # Build the docx
    print(f"\nGenerating {os.path.basename(OUTPUT_DOCX)}...")
    build_docx(entries, total)

    # Summary
    print("\n" + "=" * 60)
    print("SUMMARY")
    print("=" * 60)
    print(f"Total apps processed: {total}")
    print(f"Successful fetches:   {successes}")
    print(f"Failed fetches:       {len(failures)}")
    if failures:
        print("\nFailed package IDs:")
        for pkg_id, error in failures:
            print(f"  • {pkg_id}: {error}")
    print(f"\nOutput: {OUTPUT_DOCX}")
    print("=" * 60)


if __name__ == "__main__":
    main()
